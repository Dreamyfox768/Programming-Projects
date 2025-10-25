import tkinter as tk
from tkinter import ttk
from tkinter import filedialog
from resume_export import Exporttodoc
from not_AI_body_resume import body_part1
from body_2_resume import BodyPart2 #not ai
from personal_resume import Personal #personal info
from docx import Document

class UI:
    def __init__(self, root):
        self.document = Document()
        self.root = root
        self.root.title("🎀 Cute Resume Generator 🎀")
        self.root.geometry("800x800")
        self.root.configure(bg="#FFE6F0")  # Light pink background
        self.Exporttodoc = Exporttodoc()
        self.pi = Personal()

        self.result_var = tk.StringVar()
        self.create_widgets()

    def create_widgets(self):
        # Heading
        label = tk.Label(
            self.root, text="🎀 Cute Resume Generator 🎀",
            font=('Helvetica', 18, 'bold'), fg="#C71585",
            bg="#FFD6EC", relief='groove', bd=5, padx=10, pady=10
        )
        label.pack(pady=15)

        # Form Frame
        form_frame = tk.Frame(self.root, bg="#FFE6F0")
        form_frame.pack(pady=10, expand=True)

        for i in range(6):  # Allow columns to expand equally
            form_frame.grid_columnconfigure(i, weight=1)

        # Name Fields
        tk.Label(form_frame, text="First Name 🎀", fg="#C71585", bg="#FFE6F0").grid(row=0, column=0, sticky="e", padx=5, pady=5)
        self.first_name_entry = tk.Entry(form_frame, width=20)
        self.first_name_entry.grid(row=0, column=1, padx=5, pady=5)

        tk.Label(form_frame, text="Last Name 🎀", fg="#C71585", bg="#FFE6F0").grid(row=0, column=2, sticky="e", padx=5, pady=5)
        self.last_name_entry = tk.Entry(form_frame, width=20)
        self.last_name_entry.grid(row=0, column=3, padx=5, pady=5)

        # Email Fields
        tk.Label(form_frame, text="Email Username ✉️", fg="#C71585", bg="#FFE6F0").grid(row=1, column=0, sticky="e", padx=5, pady=5)
        self.email_entry = tk.Entry(form_frame, width=40)
        self.email_entry.grid(row=1, column=1, columnspan=3, padx=5, pady=5)

        tk.Label(form_frame, text="Email Platform 📧", fg="#C71585", bg="#FFE6F0").grid(row=2, column=0, sticky="e", padx=5, pady=5)
        self.email_platform_var = tk.StringVar(value="Gmail")
        platform_options = ["Gmail", "Yahoo", "Outlook", "iCloud", "Other"]
        platform_menu = ttk.Combobox(form_frame, textvariable=self.email_platform_var,
                                     values=platform_options, state="readonly", width=37)
        platform_menu.grid(row=2, column=1, columnspan=3, padx=5, pady=5)

        # Phone Number
        tk.Label(form_frame, text="Phone Number ☎️", fg="#C71585", bg="#FFE6F0").grid(row=3, column=0, sticky="e", padx=5, pady=5)
        self.phone_entry = tk.Entry(form_frame, width=40)
        self.phone_entry.grid(row=3, column=1, columnspan=3, padx=5, pady=5)

        # Skills and Certifications
        tk.Label(form_frame, text="Skills 💻", fg="#C71585", bg="#FFE6F0").grid(row=4, column=0, sticky="ne", padx=5, pady=5)
        self.skills_entry = tk.Text(form_frame, width=28, height=5)
        self.skills_entry.grid(row=4, column=1, padx=5, pady=5)

        tk.Label(form_frame, text="Certifications 📜", fg="#C71585", bg="#FFE6F0").grid(row=4, column=2, sticky="ne", padx=5, pady=5)
        self.certifications_entry = tk.Text(form_frame, width=28, height=5)
        self.certifications_entry.grid(row=4, column=3, padx=5, pady=5)

        # Experience
        tk.Label(form_frame, text="Job Experience 💼", fg="#C71585", bg="#FFE6F0").grid(row=5, column=0, sticky="ne", padx=5, pady=5)
        self.job_entry = tk.Text(form_frame, width=28, height=5)
        self.job_entry.grid(row=5, column=1, padx=5, pady=5)

        tk.Label(form_frame, text="Volunteer Experience 🤝", fg="#C71585", bg="#FFE6F0").grid(row=5, column=2, sticky="ne", padx=5, pady=5)
        self.volunteer_entry = tk.Text(form_frame, width=28, height=5)
        self.volunteer_entry.grid(row=5, column=3, padx=5, pady=5)

        # Result Display
        result_label = tk.Label(
            self.root, textvariable=self.result_var, bg="#FFE6F0",
            fg="#6C3483", justify="left", font=('Arial', 12), wraplength=700
        )
        result_label.pack(pady=10)

        # Buttons Frame
        btn_frame = tk.Frame(self.root, bg="#FFE6F0")
        btn_frame.pack(pady=10)

        style = {
            "width": 15,
            "fg": "white",
            "font": ("Helvetica", 10, "bold"),
            "padx": 5,
            "pady": 5
        }

        tk.Button(btn_frame, text="🎀 Submit", command=self.submit_info, bg="#FF69B4", **style).grid(row=0, column=0, padx=10)
        tk.Button(btn_frame, text="🧼 Clear", command=self.clear_fields, bg="#FFB6C1", **style).grid(row=0, column=1, padx=10)
        tk.Button(btn_frame, text="🔄 Try Again", command=self.try_again, bg="#FF1493", **style).grid(row=0, column=2, padx=10)
        tk.Button(btn_frame, text="💾 Export", command=self.export_to_file, bg="#C71585", **style).grid(row=0, column=3, padx=10)

    def submit_info(self):
        first_name = self.first_name_entry.get().strip()
        last_name = self.last_name_entry.get().strip()
        email_user = self.email_entry.get().strip()
        platform = self.email_platform_var.get()
        phone = self.phone_entry.get().strip()
        skills = self.skills_entry.get("1.0", tk.END).strip()
        certifications = self.certifications_entry.get("1.0", tk.END).strip()
        job_exp = self.job_entry.get("1.0", tk.END).strip()
        volunteer_exp = self.volunteer_entry.get("1.0", tk.END).strip()
        full_name = first_name + " " + last_name
        email = email_user + platform
        self.pi.collect_info(full_name, email, phone)
        if not first_name or not last_name or not email_user or not platform or not phone:
            self.result_var.set("⚠️ Please fill in all required fields.")
            return

        full_name = f"{first_name} {last_name}"
        email_full = f"{email_user}@{platform.lower()}.com"

        result = (
            f"🎀 Name: {full_name}\n"
            f"📧 Email: {email_full}\n"
            f"☎️ Phone: {phone}\n\n"
            f"💻 Skills:\n{skills}\n\n"
            f"📜 Certifications:\n{certifications}\n\n"
            f"💼 Job Experience:\n{job_exp}\n\n"
            f"🤝 Volunteer Experience:\n{volunteer_exp}"
        )
        self.result_var.set(result)

    def clear_fields(self):
        self.first_name_entry.delete(0, tk.END)
        self.last_name_entry.delete(0, tk.END)
        self.email_entry.delete(0, tk.END)
        self.email_platform_var.set("Gmail")
        self.phone_entry.delete(0, tk.END)
        self.skills_entry.delete("1.0", tk.END)
        self.certifications_entry.delete("1.0", tk.END)
        self.job_entry.delete("1.0", tk.END)
        self.volunteer_entry.delete("1.0", tk.END)
        self.result_var.set("")

    def try_again(self):
        self.skills_entry.delete("1.0", tk.END)
        self.certifications_entry.delete("1.0", tk.END)
        self.job_entry.delete("1.0", tk.END)
        self.volunteer_entry.delete("1.0", tk.END)
        self.result_var.set("")

    def export_to_file(self):
        resume_text = self.result_var.get()
        if not resume_text:
            self.result_var.set("⚠️ Please generate resume first using Submit.")
            return
        file_path = filedialog.asksaveasfilename(defaultextension=".txt",
                                                 filetypes=[("Text files", "*.txt")],
                                                 title="Save Resume As")
        if file_path:
            with open(file_path, "w", encoding='utf-8') as file:
                file.write(resume_text)
            self.result_var.set(f"✅ Resume exported to:\n{file_path}")
            self.build_document(file_path) #todo merge ui
    def build_document(self, file_path): #todo doc file
        self.document.add_heading("Resume", level=1)
        self.document.add_paragraph(str(self.pi)) #todo and do same for all the other
        self.document.save(file_path[:-4]+ ".docx")





def main():
    root = tk.Tk()
    app = UI(root)
    root.mainloop()


if __name__ == "__main__":
    main()
