from docx import Document
from Body_withAI import body_part1
from Body_notAI import BodyPart2
from Personal_Info import Personal

class Exporttodoc:
    """Class to build and export a resume to a Word document."""

    def __init__(self, personal: Personal, skills: BodyPart2, jobs: body_part1, file_path="resume_export.docx"):
        """
        Initialize the export class with personal info, skills, jobs, and the target file path.

        Args:
            personal (Personal): Personal information object.
            skills (BodyPart2): Skills and certifications object.
            jobs (body_part1): Job experience object.
            file_path (str, optional): Path to save the Word document. Defaults to 'resume_export.docx'.
        """
        self.document = Document()          # Create a new Word document
        self.personal_info = personal       # Store the personal info object
        self.skill_certs = skills           # Store the skills/certifications object
        self.vol_job = jobs                 # Store the job experience object
        self.file_path = file_path          # Store the output file path

    def build_resume(self):
        """Add all sections (heading, personal info, skills, jobs) to the document."""
        # Add a main heading with the full name
        self.document.add_heading(self.personal_info.full_name, level=1)

        # Add personal information section
        self.document.add_paragraph(str(self.personal_info))

        # Add skills and certifications section
        self.document.add_paragraph(str(self.skill_certs))

        # Add job experience section
        self.document.add_paragraph(str(self.vol_job))
        #add AI
        self.document.add_paragraph(str(self.vol_job.generate_job_descriptions()))


    def export(self):
        """Save the Word document to the specified file path."""
        self.document.save(self.file_path)
        print(f"Resume exported successfully as '{self.file_path}'.")
